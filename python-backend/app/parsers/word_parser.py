from __future__ import annotations

import asyncio
import re

from docx import Document

from app.parsers.base import BaseParser, ParseResult
from app.parsers.utils import normalize_file_type, truncate_text, unique_strings

HEADING_STYLE_PATTERN = re.compile(r"heading\s+(\d+)", re.IGNORECASE)


class WordParser(BaseParser):
    async def parse(self, file_path: str) -> ParseResult:
        return await asyncio.to_thread(self._parse_sync, file_path)

    def _parse_sync(self, file_path: str) -> ParseResult:
        document = Document(file_path)
        headings: list[dict[str, object]] = []
        paragraph_count = 0
        body_lines: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            level = self._extract_heading_level(paragraph.style.name if paragraph.style else "")
            if level is not None:
                headings.append({"level": level, "text": text})
                body_lines.append(f"{'#' * min(level, 6)} {text}")
            else:
                paragraph_count += 1
                body_lines.append(text)

        tables = self._extract_tables(document)
        for table in tables:
            body_lines.append(f"Table {table['table_index']}:")
            for row in table["preview_rows"]:
                body_lines.append(" | ".join(cell for cell in row if cell))

        images = self._extract_images(document)
        if images:
            body_lines.append("Images:")
            body_lines.extend(image["description"] for image in images)

        return ParseResult(
            file_type=normalize_file_type(file_path),
            summary=(
                f"Word document with {len(headings)} headings, {paragraph_count} paragraphs, "
                f"{len(tables)} tables, and {len(images)} images."
            ),
            text_content="\n".join(body_lines),
            structured_data={
                "heading_count": len(headings),
                "headings": headings[:50],
                "paragraph_count": paragraph_count,
                "table_count": len(tables),
                "tables": tables,
                "image_count": len(images),
                "images": images,
            },
            key_points=unique_strings(
                [f"Heading: {heading['text']}" for heading in headings[:3]]
                + [body_line for body_line in body_lines[:10] if body_line],
                max_items=5,
            ),
        )

    def _extract_heading_level(self, style_name: str) -> int | None:
        match = HEADING_STYLE_PATTERN.search(style_name)
        if match is None:
            return None
        return int(match.group(1))

    def _extract_tables(self, document: Document) -> list[dict[str, object]]:
        tables: list[dict[str, object]] = []

        for table_index, table in enumerate(document.tables, start=1):
            rows = [
                [cell.text.strip() for cell in row.cells]
                for row in table.rows
            ]
            preview_rows = rows[:5]
            column_count = max((len(row) for row in rows), default=0)
            tables.append(
                {
                    "table_index": table_index,
                    "row_count": len(rows),
                    "column_count": column_count,
                    "preview_rows": preview_rows,
                }
            )

        return tables

    def _extract_images(self, document: Document) -> list[dict[str, str]]:
        images: list[dict[str, str]] = []

        for image_index, inline_shape in enumerate(document.inline_shapes, start=1):
            description = ""
            inline = getattr(inline_shape, "_inline", None)
            doc_pr = getattr(inline, "docPr", None)
            if doc_pr is not None:
                description = (doc_pr.get("descr") or doc_pr.get("title") or "").strip()
            if not description:
                description = f"Image {image_index} without alt text"

            images.append(
                {
                    "image_index": str(image_index),
                    "description": truncate_text(description, 160),
                }
            )

        return images
